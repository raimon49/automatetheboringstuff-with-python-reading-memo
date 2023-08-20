#!/usr/bin/env python
# -*- coding: utf-8 -*-
# vim:fenc=utf-8 ff=unix ft=python ts=4 sw=4 sts=4 si et

def main():
    import docx

    doc = docx.Document('demo.docx')
    print(len(doc.paragraphs)) # 7

    # Documentオブジェクトが持つParagraphオブジェクトを参照
    print(doc.paragraphs[0].text) # Document Title
    print(doc.paragraphs[1].text) # A plain paragraph with some bold and some italic

    print(len(doc.paragraphs[1].runs)) # 4

    # Paragraphオブジェクトが持つRunオブジェクトを参照
    print(doc.paragraphs[1].runs[0].text) # A plain paragraph with some
    print(doc.paragraphs[1].runs[1].text) # bold
    print(doc.paragraphs[1].runs[2].text) #  and some
    print(doc.paragraphs[1].runs[3].text) # italic

    def get_text(filename):
        """ .docxファイルから全テキストを取得する関数

        Arguments:
        - filename .docxファイル名
        """
        doc = docx.Document(filename)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        return '\n\n'.join(full_text)

    print('-------- Call get_text() --------')
    print(get_text('demo.docx'))

if __name__ == '__main__':
    main()

