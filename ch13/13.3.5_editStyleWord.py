#!/usr/bin/env python
# -*- coding: utf-8 -*-
# vim:fenc=utf-8 ff=unix ft=python ts=4 sw=4 sts=4 si et

def main():
    import docx

    doc = docx.Document('demo.docx')
    print(doc.paragraphs[0].text)  # Document Title
    print(doc.paragraphs[1].style) # _ParagraphStyle('Normal') id: 139968275939920

    doc.paragraphs[0].style = 'Normal'
    print(doc.paragraphs[1].text) # A plain paragraph with some bold and some italic

    print(
        (doc.paragraphs[1].runs[0].text,
         doc.paragraphs[1].runs[1].text,
         doc.paragraphs[1].runs[2].text,
         doc.paragraphs[1].runs[3].text
         )
    ) # ('A plain paragraph with some ', 'bold', ' and some ', 'italic')


if __name__ == '__main__':
    main()

