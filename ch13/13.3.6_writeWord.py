#!/usr/bin/env python
# -*- coding: utf-8 -*-
# vim:fenc=utf-8 ff=unix ft=python ts=4 sw=4 sts=4 si et

def main():
    import docx

    doc = docx.Document()
    doc.add_paragraph('Hello world!!', 'Title')

    # doc.save('helloworld.docx')

    para_obj1 = doc.add_paragraph('これは第2段落です')
    para_obj2 = doc.add_paragraph('これは第3段落です')
    para_obj1.add_run(' これは第2段落に追加したテキストです。')

    # doc.save('multipuleParagraphs.docx')

    doc.add_heading('Header 0', 0)
    doc.add_heading('Header 1', 1)
    doc.add_heading('Header 2', 2)
    doc.add_heading('Header 3', 3)
    doc.add_heading('Header 4', 4)

    # doc.save('headings.docx')

    doc.add_paragraph('これは1ページ目')
    doc.add_page_break()
    # doc.save('twoPage.docx')

    doc.add_picture('zophie.png',
                    width=docx.shared.Inches(1),
                    height=docx.shared.Cm(4))

if __name__ == '__main__':
    main()

